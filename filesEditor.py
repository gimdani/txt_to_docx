#from hebrew_numbers import int_to_gematria
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


'''
gematia section:

Copyright (c) 2015 Ori Hoch

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.
'''


MAP = (
    (1, u'א'),
    (2, u'ב'),
    (3, u'ג'),
    (4, u'ד'),
    (5, u'ה'),
    (6, u'ו'),
    (7, u'ז'),
    (8, u'ח'),
    (9, u'ט'),
    (10, u'י'),
    (20, u'כ'),
    (30, u'ל'),
    (40, u'מ'),
    (50, u'נ'),
    (60, u'ס'),
    (70, u'ע'),
    (80, u'פ'),
    (90, u'צ'),
    (100, u'ק'),
    (200, u'ר'),
    (300, u'ש'),
    (400, u'ת'),
    (500, u'ך'),
    (600, u'ם'),
    (700, u'ן'),
    (800, u'ף'),
    (900, u'ץ')
)
separators={
    'geresh': '׳',
    'gershayim': '״'}
numerals={
    1: 'א',
    2: 'ב',
    3: 'ג',
    4: 'ד',
    5: 'ה',
    6: 'ו',
    7: 'ז',
    8: 'ח',
    9: 'ט',
    10: 'י',
    20: 'כ',
    30: 'ל',
    40: 'מ',
    50: 'נ',
    60: 'ס',
    70: 'ע',
    80: 'פ',
    90: 'צ',
    100: 'ק',
    200: 'ר',
    300: 'ש',
    400: 'ת',
    500: 'תק',
    600: 'תר',
    700: 'תש',
    800: 'תת',
    900: 'תתק'}
specials={
    0: '0',
    15: 'טו',
    16: 'טז',
    115: 'קטו',
    116: 'קטז',
    215: 'רטו',
    216: 'רטז',
    270: 'ער',
    272: 'ערב',
    274: 'עדר',
    275: 'ערה',
    298: 'רחצ',
    304: 'דש',
    315: 'שטו',
    316: 'שטז',
    344: 'שדמ',
    415: 'תטו',
    416: 'תטז',
    515: 'תקטו',
    516: 'תקטז',
    615: 'תרטו',
    616: 'תרטז',
    670: 'עתר',
    672: 'תערב',
    674: 'עדרת',
    698: 'תרחצ',
    715: 'תשטו',
    716: 'תשטז',
    744: 'תשדמ',
    815: 'תתטו',
    816: 'תתטז',
    915: 'תתקטו',
    916: 'תתקטז'}

MAP_DICT = dict([(k, v) for v, k in MAP])
GERESH = set(("'", '׳'))


def gematria_to_int(string):
    res = 0
    for i, char in enumerate(string):
        if char in GERESH and i < len(string)-1:
            res *= 1000
        if char in MAP_DICT:
            res += MAP_DICT[char]
    return res


# adapted from hebrew-special-numbers documentation
def int_to_gematria(num, gershayim=True):
    """convert integers between 1 an 999 to Hebrew numerals.
           - set gershayim flag to False to ommit gershayim
    """
    # 1. Lookup in specials
    if num in specials:
        retval = specials[num]
        return _add_gershayim(retval) if gershayim else retval

    # 2. Generate numeral normally
    parts = []
    rest = str(num)
    while rest:
        digit = int(rest[0])
        rest = rest[1:]
        if digit == 0:
            continue
        power = 10 ** len(rest)
        parts.append(numerals[power * digit])
    retval = ''.join(parts)
    # 3. Add gershayim
    return _add_gershayim(retval) if gershayim else retval


def _add_gershayim(s):
    if len(s) == 1:
        s += separators['geresh']
    else:
        s = ''.join([
            s[:-1],
            separators['gershayim'],
            s[-1:]
        ])
    return s
'''
end of gematria section
'''

if __name__ == '__main__':

   
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'David'
    font.size = Pt(18)
    font.rtl = True

    
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('normalStyle', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.rtl = True
    obj_font.size = Pt(18)
    obj_font.name = 'David'
   
    path = input("Enter file (with path!): ")
    
    file = open(path,"r+", encoding='utf8')
    #file = open(r"C:\Users\user\Downloads\book.txt","r+", encoding='utf8')

    key_words=[]
    tmp=input("""Enter keyword, tranlate (optional), heading level, and max words (optional).
you may whrite in Hebrew. for break press 'Enter'. exaples: 
Chapter
פרק
2
2
or:
הקדמה
1
or:
gate 
3:\n""")

    while True:
        if tmp =='':
            break
        val = []
        val.append(tmp)
        val.append(input())
        if gematria_to_int(val[1])==0:
            val.append(val[1])
            val[1] = val[0]
        tmp=input()
        if not tmp=='':
            val.append(tmp)
            tmp=input()
            if not tmp=='':
                val.append(tmp)
        key_words.append(val)
        tmp=input("Enter another keyword. Press Enter to escape:\n")
    print("Working on your file!")   
    book = file.readlines()
    file.close()
    for line in book:
        flag = False
        line = line.replace('\n', '') 
        for key_word in key_words:
            if key_word[0] in line and (len(key_word)<=3 or len(line.split())<=int(key_word[3])):
                line = line.replace(key_word[0], key_word[1]) 
                number = re.findall(r'\d+', line)
                if not len(number)==0:
                    line = line.replace(number[0],int_to_gematria(int(number[0])))
                h = document.add_heading(line, level=int(key_word[2]))
                h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                flag = True

        if not flag and not line=='':
            p = document.add_paragraph(line, style='normalStyle')
            #p.style = document.styles['Normal']
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save((((path.split('\\'))[-1]).split('.'))[0]+'.docx')